"""DOCX text extraction via python-docx."""
from __future__ import annotations

from io import BytesIO

from docx import Document


def extract(raw: bytes) -> str:
    doc = Document(BytesIO(raw))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)

    return "\n\n".join(parts)
