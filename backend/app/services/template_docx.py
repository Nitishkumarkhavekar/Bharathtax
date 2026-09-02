"""Officer's own .docx templates — read an uploaded Word file, and re-emit it
with fresh body content while keeping the file's LETTERHEAD (the section
header/footer: office name, DIN block, seal, signature line) exactly intact.

python-docx keeps headers/footers as separate parts referenced from the body's
final <w:sectPr>. So to "fill" a letterhead we open the original file, strip the
body paragraphs/tables (leaving the sectPr and thus the header/footer links),
and render the new content — which python-docx inserts before the sectPr.
"""
from __future__ import annotations

import io
import re

from docx import Document

from .drafting_export import render_content

_TO_RE = re.compile(r"^\s*to\s*[,:]?\s*$", re.I)
_HEADING_RE = re.compile(r"OFFICE OF|GOVERNMENT OF|F\.?\s*No|\bDIN\b", re.I)


def strip_office_heading(text: str) -> str:
    """For placing a generated draft on the officer's OWN letterhead: drop the
    draft's own office heading / DIN / F.No. / date block (everything before the
    'To,' addressee line), since the letterhead already carries the office
    identity — otherwise the heading appears twice. No-op unless the top clearly
    looks like an office heading AND a 'To,' line is found near the top."""
    lines = (text or "").splitlines()
    if not any(_HEADING_RE.search(ln) for ln in lines[:6]):
        return text
    for i, ln in enumerate(lines[:18]):
        s = ln.strip().lower()
        if _TO_RE.match(ln) or s.startswith("to,"):
            return "\n".join(lines[i:]).lstrip("\n")
    return text

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def extract_text(raw: bytes) -> str:
    """The body text of a .docx (paragraphs + tables), for the editable preview."""
    try:
        from app.ingestion.extract.docx import extract as _extract
        return (_extract(raw) or "").strip()
    except Exception:
        return ""


def _hf_has_content(hf) -> bool:
    if hf is None:
        return False
    try:
        if any((p.text or "").strip() for p in hf.paragraphs):
            return True
        xml = hf._element.xml  # type: ignore[attr-defined]
        # a logo/seal image or a table in the header/footer counts as letterhead
        return ("<w:drawing" in xml) or ("<pic:pic" in xml) or ("<w:tbl" in xml)
    except Exception:
        return False


def detect_letterhead(raw: bytes) -> bool:
    """True when any section carries header/footer content — i.e. the file has
    an office letterhead worth preserving on download."""
    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        return False
    for sec in doc.sections:
        for attr in ("header", "footer", "first_page_header", "first_page_footer",
                     "even_page_header", "even_page_footer"):
            try:
                if _hf_has_content(getattr(sec, attr, None)):
                    return True
            except Exception:
                continue
    return False


def fill_letterhead_docx(raw: bytes, content: str) -> bytes:
    """Open the officer's uploaded .docx and replace its body with `content`,
    leaving the section header/footer (letterhead) untouched. Falls back to the
    plain export if the file can't be opened."""
    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        from .drafting_export import to_docx
        return to_docx("", content)

    # Drop existing body paragraphs and tables; keep the trailing <w:sectPr>
    # (it holds the header/footer references). New paragraphs from render_content
    # are inserted before that sectPr by python-docx.
    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if tag.endswith("}p") or tag.endswith("}tbl"):
            body.remove(child)

    render_content(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
