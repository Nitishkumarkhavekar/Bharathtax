"""Render a drafted notice/order to a clean, editable .docx (Times New Roman 12,
A4, 1.5 spacing) — ready to place on the office letterhead / paste into ITBA."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

_BODY_FONT = "Times New Roman"
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
_SEP_CELL_RE = re.compile(r"^\s*:?-+:?\s*$")


def _is_table_row(s: str) -> bool:
    return bool(_TABLE_ROW_RE.match(s))


def _is_sep_row(s: str) -> bool:
    m = _TABLE_ROW_RE.match(s)
    if not m:
        return False
    cells = [c.strip() for c in m.group(1).split("|")]
    return len(cells) >= 1 and all(_SEP_CELL_RE.match(c) for c in cells)


def _add_runs(p, text: str) -> None:
    """Render **bold** segments; everything else plain."""
    for seg in re.split(r"(\*\*.+?\*\*)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**") and len(seg) > 4:
            p.add_run(seg[2:-2]).bold = True
        else:
            p.add_run(seg)


def _cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _looks_numeric(cell: str) -> bool:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", cell).replace(",", "").strip()
    if not s:
        return False
    s = re.sub(r"^(rs\.?|inr|₹|\$)\s*", "", s, flags=re.I).rstrip("%").strip()
    return bool(re.fullmatch(r"-?\d+(?:\.\d+)?", s))


def _flush_table(doc, rows_raw: list[str]) -> None:
    """Emit `rows_raw` (pipe-table lines) as a real Word table."""
    parsed: list[list[str]] = []
    for r in rows_raw:
        if _is_sep_row(r):
            continue
        m = _TABLE_ROW_RE.match(r)
        if not m:
            continue
        parsed.append([c.strip() for c in m.group(1).split("|")])
    if not parsed:
        return
    n_cols = max(len(r) for r in parsed)
    parsed = [r + [""] * (n_cols - len(r)) for r in parsed]
    body = parsed[1:]
    right_cols: set[int] = set()
    for c in range(n_cols):
        col = [row[c] for row in body if row[c].strip()]
        if col and sum(_looks_numeric(x) for x in col) / len(col) >= 0.6:
            right_cols.add(c)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    for i, row in enumerate(parsed):
        for j, raw in enumerate(row):
            cell = table.cell(i, j)
            _cell_borders(cell)
            p = cell.paragraphs[0]
            p.text = ""
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if j in right_cols else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Body rows render bold cells via inline `**...**`; header row bolds
            # every cell.
            if i == 0:
                p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", raw)).bold = True
            else:
                _add_runs(p, raw)
    doc.add_paragraph()


def to_docx(title: str, content: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2.5)
        section.left_margin = section.right_margin = Cm(2.5)

    table_buf: list[str] = []
    for raw in (content or "").splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        # Accumulate pipe-table rows.
        if _is_table_row(stripped):
            table_buf.append(stripped)
            continue
        # First non-table line — commit the table before continuing.
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []
        # blank line -> a spacer paragraph
        if not stripped:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        # A markdown heading (#...) or a fully-bold short line -> bold, no marker.
        m = re.match(r"^\s{0,3}#{1,6}\s+(.*)$", line)
        if m:
            p.add_run(m.group(1).strip()).bold = True
        else:
            _add_runs(p, line)
    if table_buf:
        _flush_table(doc, table_buf)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
