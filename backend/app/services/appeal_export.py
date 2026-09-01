"""Render the draft appellate order to an editable .docx in the standard
CIT(A) / NFAC layout the user expects:

  * Centred, bold + underlined title — "<NAME> (PAN:<PAN>) A.Y. <YEAR>"
  * Times New Roman 12 pt body, 1.5 line spacing, justified
  * Numbered paragraphs use a TAB after the number ("1.\\t…") with a hanging
    indent so the second line of a paragraph aligns under the first character
    of the body, not under the number — matches the screenshot the user shared.
  * Sub-numbering (2.1, (a), (b), (i)) gets its own deeper indent.
  * Bold inline markdown (**…**) is preserved as bold runs.
  * Section labels the LLM emits as ALL-CAPS lines or `**HEADING**` get
    centred + bold.

The function accepts either the legacy (title, draft_text) call OR a full
keyword-form with the case metadata so the header can be built deterministically.
"""
from __future__ import annotations

import io
import re

from docx import Document as Docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_NUM_RE = re.compile(r"^(\d+(?:\.\d+)?)\s*[.)]\s+(.*)$")          # 1. / 2.1 …
_LETTER_RE = re.compile(r"^\(([a-z]|[ivxlcdm]+)\)\s+(.*)$", re.I)  # (a), (ii)…
_BULLET_RE = re.compile(r"^\s*[-*•]\s+(.*)$")
# A markdown pipe-table row: starts with `|`, contains at least one internal
# `|`, ends with `|`. We match liberally on internal whitespace so the LLM's
# alignment styles (`| :--- | ---: |` etc.) still count as rows.
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
# A markdown-table SEPARATOR row (the `|:---|---:|:---:|` line that follows
# the header) — cells that are only dashes/colons/spaces. Detected once we
# know the line is already a table row.
_TABLE_SEP_CELL_RE = re.compile(r"^\s*:?-+:?\s*$")
# Sentinel prefix marking a pre-assembled pipe-table block so the render loop
# knows to route it through _render_markdown_table instead of _render_line.
_TABLE_BLOCK_MARK = "\x00TABLE\x00"


def _is_table_row(s: str) -> bool:
    return bool(_TABLE_ROW_RE.match(s))


def _is_table_separator(s: str) -> bool:
    m = _TABLE_ROW_RE.match(s)
    if not m:
        return False
    cells = [c.strip() for c in m.group(1).split("|")]
    return len(cells) >= 1 and all(_TABLE_SEP_CELL_RE.match(c) for c in cells)


def _parse_table_rows(rows: list[str]) -> list[list[str]]:
    """Take raw pipe-table lines and return a matrix of cells (strings).
    The alignment-separator row (`| :--- | ---: |`) is dropped."""
    out: list[list[str]] = []
    for r in rows:
        if _is_table_separator(r):
            continue
        m = _TABLE_ROW_RE.match(r)
        if not m:
            continue
        cells = [c.strip() for c in m.group(1).split("|")]
        out.append(cells)
    return out


def _looks_numeric(cell: str) -> bool:
    """Right-align a cell when it reads as an amount / count. Strips markdown
    bold markers, Indian-style comma-grouped digits, currency prefix, plus
    optional trailing `-` for negatives."""
    s = _strip_bold_markers(cell).replace(",", "").strip()
    if not s:
        return False
    # accept: `1,60,500`, `Rs. 63,94,900`, `63,94,900.00`, `-15,000`, `10%`
    s = re.sub(r"^(rs\.?|inr|₹|\$)\s*", "", s, flags=re.I)
    s = s.rstrip("%").strip()
    return bool(re.fullmatch(r"-?\d+(?:\.\d+)?", s))

BODY_FONT = "Times New Roman"
BODY_PT = 12


# --------------------------------------------------------------- helpers
def _strip_bold_markers(s: str) -> str:
    return _BOLD_RE.sub(r"\1", s).strip()


def _add_runs_with_bold(p, text: str, *, bold_all: bool = False,
                       italic_all: bool = False) -> None:
    """Append `text` to paragraph `p`, parsing `**bold**` inline markdown.

    `italic_all` forces every run (bold and normal) into italic, so the paragraph
    reads as an italicised legal-submission block while still allowing embedded
    `**strong**` emphasis inside it.
    """
    pos = 0
    any_added = False
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            r.bold = bold_all
            r.italic = italic_all
            any_added = True
        r = p.add_run(m.group(1))
        # Body text must not be bold — the LLM often wraps whole sentences in
        # **...**; render those inline as normal weight (headings are rendered
        # separately with explicit bold) so the order reads like a proper legal
        # document (bold headings, plain justified body).
        r.bold = bold_all
        r.italic = italic_all
        any_added = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        r.bold = bold_all
        r.italic = italic_all
        any_added = True
    # Edge case: empty text — keep paragraph empty (no runs).
    if not any_added:
        return


def _apply_body_font(p) -> None:
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.size = Pt(BODY_PT)
        # East-Asia + complex-script fallbacks so the docx opens consistently.
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), BODY_FONT)
        rFonts.set(qn("w:hAnsi"), BODY_FONT)
        rFonts.set(qn("w:cs"), BODY_FONT)


def _hanging(p, *, first_left_cm: float, hang_cm: float) -> None:
    """Hanging indent: left margin sits at `first_left_cm`, the tab stop /
    continuation lines align at `first_left_cm + hang_cm`."""
    pf = p.paragraph_format
    pf.left_indent = Cm(first_left_cm + hang_cm)
    pf.first_line_indent = Cm(-hang_cm)
    # A tab stop at the body alignment so the visible TAB after the number
    # lands on it (matches the screenshot's "1.\t" gap).
    tabs = pf.tab_stops
    try:
        tabs.add_tab_stop(Cm(first_left_cm + hang_cm))
    except ValueError:
        pass


def _body_paragraph(doc, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    space_after_pt: float = 6.0, line_spacing: float = 1.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    return p


def _underline_run(r) -> None:
    rPr = r._element.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)


def _set_doc_defaults(doc) -> None:
    """Document-wide font, page margins, and an explicit default paragraph
    style so any paragraph we *don't* touch still uses Times New Roman 12."""
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_PT)
    # Keep the Normal style upright by default; italic is applied per-run so
    # the export can mix upright headings with italic body without style
    # inheritance turning everything italic.
    style.font.italic = False
    # Force the eastAsia + complex-script names too.
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    # Page setup: A4 margins 2.5 cm.
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# --------------------------------------------------------------- title block
def _format_title(case_title: str, *, pan: str | None, ay: str | None) -> str:
    """Build the centred title line — e.g.
       'RAJU (PAN:ALKPR4293F) A.Y. 2017-18'.

    We extract the appellant name from the case title (drop trailing 'case').
    """
    name = (case_title or "").strip()
    name = re.sub(r"\s+case$", "", name, flags=re.I).strip()
    parts = [name.upper()]
    if pan:
        parts.append(f"(PAN:{pan.strip().upper()})")
    if ay:
        # Accept 'AY 2017-18', '2017-18', 'A.Y. 2017-18' — normalise to A.Y.
        ay_clean = re.sub(r"^a\.?y\.?\s*", "", ay.strip(), flags=re.I).strip()
        parts.append(f"A.Y. {ay_clean}")
    return " ".join(p for p in parts if p)


def _add_title(doc, line: str) -> None:
    if not line.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(line)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(13)
    _underline_run(r)


# --------------------------------------------------------------- body rendering
def _looks_like_short_label(stripped: str) -> bool:
    """A `**...**`-wrapped paragraph is only a heading if it's SHORT and
    doesn't end like a sentence — otherwise it's a bold body statement
    (e.g. a fully-bold submission line) which should render inline, not
    as a centred banner."""
    if not (stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4):
        return False
    body = stripped[2:-2].strip()
    if not body:
        return False
    # Sentence-shaped → NOT a heading.
    if body.endswith((".", ";", ",", "?", "!")):
        return False
    if len(body) > 90:
        return False
    if len(body.split()) > 12:
        return False
    return True


def _is_heading_line(stripped: str) -> bool:
    """Detect section-header lines the LLM emits (INTRODUCTION, GROUNDS OF
    APPEAL, DISCUSSION AND FINDINGS, etc.). These stay upright (non-italic).

    We are DELIBERATELY conservative here — the previous heuristic classified
    any short ALL-CAPS-ish line as a heading, which caught body sentences with
    statute references (e.g. "S.144/S.147 ORDER DATED 07.12.2019") and turned
    them into centred banners.
    """
    if _looks_like_short_label(stripped):
        return True
    if re.match(r"^#{1,4}\s+", stripped):
        return True
    # ALL-CAPS heading fallback — require:
    #  * no sentence-ending punctuation
    #  * no lower-case letters at all
    #  * at least two alphabetic words
    #  * short line (≤ 60 chars) so a body sentence with a couple of
    #    all-caps statute refs doesn't trip it
    if stripped.endswith((".", ",", ";", ":", "?", "!")):
        return False
    if len(stripped) > 60:
        return False
    if any(c.islower() for c in stripped):
        return False
    if not re.fullmatch(r"[A-Z0-9 ,&/\-:.()]+", stripped):
        return False
    if len([w for w in stripped.split() if any(c.isalpha() for c in w)]) < 2:
        return False
    return True


def _render_line(doc, line: str) -> None:
    """Render a single non-empty source line into the right paragraph kind.

    The house style — matching the reference screenshots — is:
      * Title: centred bold + underlined, upright (non-italic).
      * Section headings: centred bold, upright.
      * Body paragraphs (numbered, lettered, plain): Times New Roman 12pt
        italic, justified, 1.5 line spacing, with a 1 cm hanging indent so
        the number/letter column sits on the left and continuation lines
        wrap under the first character of the body.
    """
    stripped = line.strip()

    # ---- Section headings — bold + centred + UPRIGHT (but only for lines
    # that actually LOOK like a heading, not sentence-shaped bold bodies) ----
    if _looks_like_short_label(stripped):
        p = _body_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after_pt=10)
        r = p.add_run(_strip_bold_markers(stripped))
        r.bold = True
        r.italic = False
        _apply_body_font(p)
        return
    h = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if h:
        text = _strip_bold_markers(h.group(2))
        # `####` (h4) is our sub-heading marker → left-aligned bold, upright,
        # with a bit of extra top space so the reader can find it. `#`–`###`
        # remain top-of-section banners → centred.
        alignment = (WD_ALIGN_PARAGRAPH.LEFT if h.group(1) == "####"
                     else WD_ALIGN_PARAGRAPH.CENTER)
        p = _body_paragraph(doc, alignment=alignment, space_after_pt=8)
        r = p.add_run(text)
        r.bold = True
        r.italic = False
        _apply_body_font(p)
        return
    if _is_heading_line(stripped):
        p = _body_paragraph(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after_pt=10)
        r = p.add_run(stripped)
        r.bold = True
        r.italic = False
        _apply_body_font(p)
        return

    # ---- Numbered top-level paragraph: "1.  text" or "2.1  text" (italic) ----
    m = _NUM_RE.match(stripped)
    if m:
        num = m.group(1)
        body = m.group(2)
        first_left = 0.0
        hang = 1.0  # 1 cm hanging indent
        p = _body_paragraph(doc)
        _hanging(p, first_left_cm=first_left, hang_cm=hang)
        r_num = p.add_run(f"{num}.")
        r_num.italic = True
        r_num.bold = True    # the numbering itself reads a touch heavier
        p.add_run().add_tab()
        _add_runs_with_bold(p, body, italic_all=True)
        _apply_body_font(p)
        return

    # ---- Lettered / roman sub-ground: "(a) text", "(i) text" (italic) ----
    m = _LETTER_RE.match(stripped)
    if m:
        marker = m.group(1)
        body = m.group(2)
        first_left = 1.5
        hang = 1.0
        p = _body_paragraph(doc)
        _hanging(p, first_left_cm=first_left, hang_cm=hang)
        r_marker = p.add_run(f"({marker})")
        r_marker.italic = True
        p.add_run().add_tab()
        _add_runs_with_bold(p, body, italic_all=True)
        _apply_body_font(p)
        return

    # ---- Bullet (italic) ----
    m = _BULLET_RE.match(stripped)
    if m:
        body = m.group(1)
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        _add_runs_with_bold(p, body, italic_all=True)
        _apply_body_font(p)
        return

    # ---- Plain paragraph (italic) ----
    p = _body_paragraph(doc)
    _add_runs_with_bold(p, stripped, italic_all=True)
    _apply_body_font(p)


def _set_cell_borders(cell) -> None:
    """Apply thin single-line borders on all four sides of a table cell so
    the exported table doesn't look like plain columns of text. Uses raw
    OXML because python-docx doesn't expose per-cell borders directly."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 1/2 pt line
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _render_markdown_table(doc, block: str) -> None:
    """Render a stripped markdown-pipe-table block as a real Word table with
    a bold header row, thin grey borders and right-aligned numeric columns.
    `block` is the text produced by _collapse_lines with the sentinel
    prefix stripped — each line is one raw pipe row (separator dropped)."""
    rows = _parse_table_rows(block.split("\n"))
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad ragged rows so add_table doesn't crash on mismatched cell counts.
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    # Decide right-alignment per column by inspecting body cells (skip header).
    body = rows[1:] if len(rows) > 1 else []
    right_cols: set[int] = set()
    if body:
        for c in range(n_cols):
            col_cells = [row[c] for row in body if row[c].strip()]
            if col_cells and sum(_looks_numeric(x) for x in col_cells) / len(col_cells) >= 0.6:
                right_cols.add(c)

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    # Give the whole table a small top gap so it doesn't butt into the
    # previous paragraph.
    for i, row in enumerate(rows):
        for j, raw in enumerate(row):
            cell = table.cell(i, j)
            _set_cell_borders(cell)
            # Clear the default empty paragraph docx inserts, then write ours.
            p = cell.paragraphs[0]
            p.text = ""
            if j in right_cols and i > 0:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif j in right_cols and i == 0:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Header row = bold + upright; body = upright non-italic + tabular
            # numerals for numeric columns.
            _add_runs_with_bold(p, raw, bold_all=(i == 0), italic_all=False)
            for r in p.runs:
                r.font.name = BODY_FONT
                r.font.size = Pt(BODY_PT)
    # A trailing empty paragraph so the following text isn't glued to the
    # bottom border.
    doc.add_paragraph()


def _normalize_asterisks(text: str) -> str:
    """Clean stray asterisks the LLM sometimes emits (e.g. "****Result:"):
    collapse runs of 3+ to a pair, and drop unbalanced ** on a line so no raw
    asterisks render in the order. Balanced **bold** is preserved."""
    out = []
    for ln in (text or "").split("\n"):
        s = re.sub(r"\*{3,}", "**", ln)
        if len(re.findall(r"\*\*", s)) % 2 != 0:
            s = s.replace("**", "")
        out.append(s)
    return "\n".join(out)


def _collapse_lines(draft_text: str) -> list[str]:
    """Group wrapped lines together so a single logical paragraph (which the
    LLM may have wrapped at ~80 chars) renders as one Word paragraph.
    Consecutive markdown pipe-table rows are collected into ONE special block
    prefixed with `_TABLE_BLOCK_MARK` so the render loop can dispatch to
    the real Word-table renderer instead of rendering them as prose."""
    raw_lines = (draft_text or "").replace("\r\n", "\n").split("\n")
    blocks: list[str] = []
    buf: list[str] = []
    table_buf: list[str] = []

    def flush() -> None:
        if buf:
            blocks.append(" ".join(buf).strip())
            buf.clear()

    def flush_table() -> None:
        if table_buf:
            blocks.append(_TABLE_BLOCK_MARK + "\n".join(table_buf))
            table_buf.clear()

    for ln in raw_lines:
        s = ln.rstrip()
        if not s.strip():
            flush()
            flush_table()
            blocks.append("")  # explicit blank between paragraphs
            continue
        stripped = s.strip()
        # Markdown pipe-table row — accumulate contiguously, don't merge into
        # the prose paragraph buffer.
        if _is_table_row(stripped):
            flush()
            table_buf.append(stripped)
            continue
        # First non-table line after a table run — commit the table before
        # continuing with prose.
        if table_buf:
            flush_table()
        # A markdown heading (#..####) or a fully-**bold** label line is a
        # STANDALONE block: flush the previous paragraph AND emit the heading on
        # its own, so the body line that follows is not merged into (and bolded
        # as part of) the heading.
        is_heading = bool(re.match(r"^#{1,4}\s", stripped)) or (
            stripped.startswith("**") and stripped.endswith("**"))
        if is_heading:
            flush()
            blocks.append(stripped)
            continue
        # Other structural starts (numbered / lettered / bullet) begin a new
        # paragraph but may still wrap onto following lines.
        starts_new = bool(
            _NUM_RE.match(stripped)
            or _LETTER_RE.match(stripped)
            or _BULLET_RE.match(stripped)
        )
        if starts_new and buf:
            flush()
        buf.append(stripped)
    flush()
    flush_table()
    # Drop leading blanks and double blanks.
    out: list[str] = []
    last_blank = False
    for b in blocks:
        if not b:
            if not last_blank and out:
                out.append("")
            last_blank = True
        else:
            out.append(b)
            last_blank = False
    return out


# --------------------------------------------------------------- public API
def build_order_docx(
    case_title: str,
    draft_text: str,
    *,
    pan: str | None = None,
    assessment_year: str | None = None,
    section: str | None = None,
    banner_text: str = "DRAFT APPELLATE ORDER",
    subtitle_prefix: str = "Appeal against order u/s",
    footer_text: str = (
        "Draft generated by BharatTax for the Commissioner (Appeals) to "
        "apply independent mind and finalise. Verify every citation against "
        "the cited source."
    ),
) -> bytes:
    """Render the draft into a CIT(A)/NFAC-style .docx.

    Backwards-compatible signature: callers that only pass (title, text)
    still work; the title block then uses just the case title. The AO-side
    assessment engine passes its own banner / subtitle / footer so the same
    house-style renderer produces an assessment order.
    """
    doc = Docx()
    _set_doc_defaults(doc)

    # Optional small banner above the centred title.
    if banner_text:
        banner = doc.add_paragraph()
        banner.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        banner.paragraph_format.space_after = Pt(4)
        br = banner.add_run(banner_text)
        br.bold = True
        br.font.name = BODY_FONT
        br.font.size = Pt(10)

    title_line = _format_title(case_title, pan=pan, ay=assessment_year)
    _add_title(doc, title_line)

    if section and subtitle_prefix:
        sub = doc.add_paragraph()
        sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(14)
        r = sub.add_run(f"{subtitle_prefix} {section.strip()}")
        r.italic = True
        r.font.name = BODY_FONT
        r.font.size = Pt(11)

    for block in _collapse_lines(_normalize_asterisks(draft_text)):
        if not block:
            # Blank separator: a small gap is already provided by space_after
            # on each paragraph — only add an explicit empty paragraph if the
            # previous block was a centred heading.
            doc.add_paragraph()
            continue
        if block.startswith(_TABLE_BLOCK_MARK):
            _render_markdown_table(doc, block[len(_TABLE_BLOCK_MARK):])
            continue
        _render_line(doc, block)

    # Footer disclaimer
    note = doc.add_paragraph()
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run(footer_text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = BODY_FONT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
