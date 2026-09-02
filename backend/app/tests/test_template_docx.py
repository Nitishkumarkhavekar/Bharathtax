"""Officer's own .docx templates — letterhead (header/footer) must survive a
body re-fill."""
import io

from docx import Document

from app.services import template_docx as td


def _doc_with_letterhead() -> bytes:
    d = Document()
    d.sections[0].header.paragraphs[0].text = "OFFICE OF THE ASSESSING OFFICER, WARD 2(3), PUNE"
    d.sections[0].footer.paragraphs[0].text = "DIN 2026 · signature"
    d.add_paragraph("Original body one")
    d.add_paragraph("Original body two")
    b = io.BytesIO(); d.save(b)
    return b.getvalue()


def test_detect_letterhead():
    assert td.detect_letterhead(_doc_with_letterhead()) is True
    # a plain doc with no header/footer text
    d = Document(); d.add_paragraph("just a body")
    b = io.BytesIO(); d.save(b)
    assert td.detect_letterhead(b.getvalue()) is False


def test_extract_text():
    txt = td.extract_text(_doc_with_letterhead())
    assert "Original body one" in txt and "Original body two" in txt


def test_fill_preserves_header_footer_and_replaces_body():
    out = td.fill_letterhead_docx(_doc_with_letterhead(), "New content\n\n**Heading**\nline two")
    d = Document(io.BytesIO(out))
    # letterhead intact
    assert "ASSESSING OFFICER" in d.sections[0].header.paragraphs[0].text
    assert "DIN 2026" in d.sections[0].footer.paragraphs[0].text
    # new body present, old body gone
    bodies = [p.text for p in d.paragraphs if p.text.strip()]
    assert "New content" in bodies and "Heading" in bodies and "line two" in bodies
    assert all("Original body" not in b for b in bodies)


def test_fill_on_garbage_falls_back():
    # not a real .docx → still returns a usable .docx (plain export), no crash
    out = td.fill_letterhead_docx(b"not a docx", "hello")
    d = Document(io.BytesIO(out))
    assert any("hello" in p.text for p in d.paragraphs)
