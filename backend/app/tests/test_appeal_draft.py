from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument

from app.models.appeal import AppealCase, AppealDocument
from app.services import appeal_draft
from app.ingestion.extract import extract_text


def _build_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Ground No. 1: The rectification order is bad in law.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "PAN"
    table.rows[0].cells[1].text = "ALKPR4293F"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_text_supports_docx():
    raw = _build_docx_bytes()
    text = extract_text(
        raw,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="appeal.docx",
    )
    assert "Ground No. 1" in text
    assert "ALKPR4293F" in text


def test_docs_text_includes_inventory_and_all_documents():
    case = AppealCase(title="Raju", owner_user_id=1, wing_id=1)
    case.documents = [
        AppealDocument(
            filename="FORM35.pdf",
            category="form_35",
            text="Form 35 appeal details and verification.",
            pages=3,
            minio_key="x",
        ),
        AppealDocument(
            filename="submission.docx",
            category="written_submission",
            text="Written submission explaining why the rectification order should be cancelled.",
            pages=0,
            minio_key="y",
        ),
    ]

    bundle = appeal_draft._docs_text(case, max_chars=8000)

    assert "DOCUMENT INVENTORY" in bundle
    assert "FORM35.pdf" in bundle
    assert "submission.docx" in bundle
    assert "Written submission" in bundle


def test_issue_doc_context_prioritises_relevant_documents():
    case = AppealCase(title="Raju", owner_user_id=1, wing_id=1)
    case.documents = [
        AppealDocument(
            filename="rectification_order.pdf",
            category="assessment_order",
            text="This rectification order under section 154 raised a fresh demand.",
            pages=2,
            minio_key="a",
        ),
        AppealDocument(
            filename="postal_tracking.pdf",
            category="unclassified",
            text="India Post tracking details only.",
            pages=1,
            minio_key="b",
        ),
    ]

    context = appeal_draft._issue_doc_context(case, "Validity of rectification order under section 154")

    assert "rectification_order.pdf" in context
    assert "section 154" in context
