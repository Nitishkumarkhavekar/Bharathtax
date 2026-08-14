"""Generate BharathTax technical design documentation as a .docx file.

Run:  python scripts/generate_docs.py
Output: BharathTax_Documentation.docx (at project root)
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


# ---------- styling helpers ----------

def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F4F4")
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        c1.text = k
        c2.text = v
        for run in c1.paragraphs[0].runs:
            run.bold = True


def add_table_from_grid(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
        shade_cell(hdr[i], "D9E2F3")
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.cell(r, c).text = val


def set_doc_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name, size in [("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.font.bold = True


# ---------- document body ----------

def build_doc(out_path: Path) -> None:
    doc = Document()
    set_doc_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # ---------- COVER ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BharathTax")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Technical Design & Architecture Documentation")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(
        "Citation-grounded Indian Income Tax research assistant\n"
        "Self-hosted · Retrieval-Augmented Generation (RAG) · Anti-hallucination guarantee"
    )
    r.font.size = Pt(12)
    r.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Document Version 1.0  ·  June 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ---------- TABLE OF CONTENTS (manual) ----------
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Project Overview",
        "3. Source Corpus",
        "4. Dataset Design Philosophy",
        "5. Chunk Schema — The Unit of the Dataset",
        "6. Ingestion Pipeline",
        "7. Retrieval & Query Flow",
        "8. Chat History & Conversational Memory",
        "9. Evaluation Methodology",
        "10. Token Efficiency at Inference",
        "11. Hardware Requirements",
        "12. Deployment Patterns",
        "13. Performance Optimization",
        "14. Security, Compliance & Refusal Logic",
        "15. Future Roadmap",
        "16. Appendix — Example Chunk Record",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ---------- 1. EXECUTIVE SUMMARY ----------
    doc.add_heading("1. Executive Summary", level=1)
    add_para(
        doc,
        "BharathTax is a self-hosted, citation-grounded research assistant for Indian "
        "income tax law. It answers user questions using only primary statutory sources "
        "(Income Tax Act, Income Tax Rules, and CBDT circulars), reranks retrieved "
        "passages with a cross-encoder, and produces answers that always cite the exact "
        "Section or Rule used. When the corpus does not contain a relevant answer, the "
        "system refuses rather than hallucinates — this anti-hallucination guarantee is "
        "the central design constraint of the product."
    )
    add_para(
        doc,
        "This document captures the architectural reasoning behind the system: how the "
        "dataset is prepared from PDF sources, how chunks and embeddings are designed, "
        "how queries flow through retrieval and generation, how to evaluate quality "
        "honestly, how to manage tokens efficiently at inference, what hardware is "
        "needed to serve traffic at scale, and how to split the deployment across "
        "multiple virtual machines."
    )

    # ---------- 2. PROJECT OVERVIEW ----------
    doc.add_heading("2. Project Overview", level=1)

    doc.add_heading("2.1 What it is", level=2)
    add_para(
        doc,
        "A self-hosted Retrieval-Augmented Generation (RAG) system that turns Indian "
        "tax statutes into an interactive question-answering assistant. The product is "
        "multi-domain by design (Income Tax today; GST and Customs are configuration-"
        "only additions). Access is gated through a department / wing / seat licence "
        "model suitable for government and enterprise deployments."
    )

    doc.add_heading("2.2 Two workstreams, one database", level=2)
    add_bullet(doc, "Authentication, licensing, hybrid retrieval, Ask Bot, document Q&A, admin, audit logs.", bold_lead="Workstream A — Application:")
    add_bullet(doc, "Configuration-driven, idempotent, re-runnable pipeline that turns law PDFs into a chunked, embedded, indexed corpus.", bold_lead="Workstream B — Ingestion:")

    doc.add_heading("2.3 Service architecture", level=2)
    add_para(
        doc,
        "docker compose runs the following services: postgres (with pgvector extension), "
        "redis (Celery broker), minio (raw file object store), ml-server (FastAPI "
        "wrapping bge-m3 embeddings and bge-reranker-v2-m3 cross-encoder), api (FastAPI "
        "application), worker and beat (Celery for ingest and scheduled jobs), and "
        "frontend (React + Vite). The LLM is not a containerised service in development "
        "— it sits behind an LLMClient interface with three implementations: mock (no "
        "model), ollama (development), and vllm (production). The implementation is "
        "selected purely through environment variables, with no code change."
    )

    doc.add_heading("2.4 Request flow at a glance", level=2)
    add_code_block(doc, (
        "frontend  →  api /ask  →  retrieval (dense pgvector + sparse Postgres FTS\n"
        "                                     → merge → cross-encoder rerank\n"
        "                                     → parent-section expansion)\n"
        "                       →  grounding gate (refusal logic)\n"
        "                       →  LLM\n"
        "                       →  answer + structured citations"
    ))

    # ---------- 3. SOURCE CORPUS ----------
    doc.add_heading("3. Source Corpus", level=1)
    add_para(
        doc,
        "The seed corpus consists of four primary-source PDFs, totalling approximately "
        "458 MB, all sourced from official Indian government publications:"
    )
    add_table_from_grid(doc, ["Document", "Approx. Size", "Role"], [
        ["Income-tax Act 1961", "~180 MB", "Primary statute (in force through FY 2025-26)"],
        ["Income-tax Act 2025", "~107 MB", "Successor statute (effective FY 2026-27)"],
        ["Income-tax Rules 1962", "~104 MB", "Subordinate legislation under Act 1961"],
        ["Income-tax Rules 2026", "~72 MB",  "Subordinate legislation under Act 2025"],
    ])
    add_para(
        doc,
        "The presence of both 1961 and 2025 Acts in parallel is intentional and "
        "non-negotiable: queries about prior assessment years must resolve against the "
        "1961 framework, while AY 2026-27 onward must resolve against the 2025 framework. "
        "This requires per-chunk effective-date metadata (see §5)."
    )

    # ---------- 4. DATASET DESIGN PHILOSOPHY ----------
    doc.add_heading("4. Dataset Design Philosophy", level=1)

    doc.add_heading("4.1 Why RAG, not fine-tuning", level=2)
    add_para(
        doc,
        "A common but incorrect instinct is to build a Q&A dataset and fine-tune a "
        "language model on it. For Indian tax law specifically, three reasons make "
        "this approach unsuitable:"
    )
    add_bullet(doc, "A wrong section number is a wrong tax filing. Users must be able to verify each fact against the cited statute. Fine-tuned models do not produce verifiable citations.", bold_lead="Hallucination is catastrophic.")
    add_bullet(doc, "Tax law changes every fiscal year through the Finance Act, and continuously through CBDT circulars and notifications. A fine-tuned model rots; a RAG index simply re-ingests.", bold_lead="The law changes constantly.")
    add_bullet(doc, "It is not possible to enumerate every taxpayer question in advance. Retrieval over the actual statute handles the long tail of unseen questions.", bold_lead="Coverage is impossible.")

    doc.add_heading("4.2 Why structure-aware chunking", level=2)
    add_para(
        doc,
        "Indian statutory drafting is deeply hierarchical: Chapter → Section → "
        "sub-section → clause → proviso → Explanation. Blind fixed-size windowing "
        "(e.g., splitting every 512 tokens) destroys this structure — a proviso may be "
        "split in half, an Explanation may be detached from the rule it modifies, and "
        "the resulting chunks cannot carry a precise citation. Structure-aware "
        "chunking instead respects the legal tree: one parent chunk per Section "
        "(retaining full context for the LLM) and one child chunk per sub-section / "
        "clause / proviso / Explanation (used for precise retrieval matching)."
    )

    doc.add_heading("4.3 The breadcrumb principle", level=2)
    add_para(
        doc,
        "Every chunk is prefixed with its legal-path breadcrumb, e.g. "
        "“Income Tax Act 1961 > Chapter VIA > Section 80C > sub-section (2) > clause (a)”. "
        "This prefix is part of the text that gets embedded — so the embedding captures "
        "not only the semantic content but also the structural location. A query "
        "mentioning “80C” will match the correct chunk even when the body text does "
        "not explicitly contain the number 80C."
    )

    # ---------- 5. CHUNK SCHEMA ----------
    doc.add_heading("5. Chunk Schema — The Unit of the Dataset", level=1)
    add_para(
        doc,
        "Each row in the chunks table represents one retrieval / citation unit. The "
        "schema is intentionally rich: structural fields are stored as first-class "
        "columns (not buried in prose), which enables both precise filtering and "
        "precise citation."
    )

    add_table_from_grid(doc, ["Field", "Type", "Purpose"], [
        ["chunk_id", "text PK", "Stable identifier, e.g. ita1961_s80E_ss1"],
        ["text", "text", "Breadcrumb-prefixed body — this is what gets embedded"],
        ["body", "text", "Raw body without the breadcrumb (cleaner LLM input)"],
        ["breadcrumb", "text", "Human-readable legal path"],
        ["level", "enum", "section | subsection | clause | proviso | explanation | rule"],
        ["act_name", "text", "Income Tax Act 1961 / 2025, etc."],
        ["chapter", "text", "Chapter number (Roman or Arabic)"],
        ["section_number", "text", "e.g. 80C, 80E"],
        ["subsection", "text", "e.g. (1), (2)"],
        ["clause", "text", "e.g. (a), (b)"],
        ["proviso_no", "int", "Ordinal of the proviso within the parent unit"],
        ["explanation_no", "int", "Ordinal of the Explanation within the parent unit"],
        ["rule_number", "text", "For Rules ingestion"],
        ["effective_date", "date", "Non-negotiable. Drives in-force filtering"],
        ["status", "enum", "in_force | repealed | amended"],
        ["amended_by", "jsonb", "List of Finance Acts that touched this provision"],
        ["cross_references", "jsonb", "Other sections cited from within this chunk"],
        ["content_type", "enum", "prose | rate_table | form_schema"],
        ["parent_chunk_id", "text FK", "Narrow child → full parent for LLM context"],
        ["source_doc", "text", "Filename of the originating PDF"],
        ["source_page", "int", "Page number for clickable citations"],
        ["source_checksum", "text", "SHA-256 of the source file (idempotency)"],
        ["embedding", "vector(1024)", "bge-m3 embedding of the `text` field"],
    ])

    doc.add_heading("5.1 A worked example — Section 80E, sub-section (1)", level=2)
    add_para(
        doc,
        "The single record below represents how one sub-section of the Income Tax Act "
        "1961 is prepared and stored. Section 80E (deduction for interest on higher-"
        "education loan) is a useful illustrative example because it exercises every "
        "field — it has structural depth (section → sub-section → clause), an "
        "Explanation, cross-references to other sections, and known amendments."
    )
    add_code_block(doc, EXAMPLE_CHUNK_JSON)

    doc.add_heading("5.2 Why each field exists", level=2)
    add_bullet(doc, "The text field carries the breadcrumb prefix so the embedding captures both meaning and legal location. Big retrieval-quality win.", bold_lead="text (breadcrumb-prefixed):")
    add_bullet(doc, "The LLM sees this in the prompt — cleaner reading without the breadcrumb noise.", bold_lead="body:")
    add_bullet(doc, "Without these, Act 1961 and Act 2025 surface together and the chatbot confidently mixes repealed and current law. Mandatory filters at retrieval time.", bold_lead="effective_date / status:")
    add_bullet(doc, "Enables “has this section been amended recently?” questions without a separate corpus.", bold_lead="amended_by:")
    add_bullet(doc, "At retrieval time, the system can expand related sections (e.g. retrieve 80E together with 10(23C) it references).", bold_lead="cross_references:")
    add_bullet(doc, "Retrieve narrow child for precision; substitute the full parent section for LLM context.", bold_lead="parent_chunk_id:")
    add_bullet(doc, "When the user clicks the citation in the UI, they jump to the source PDF page for verification.", bold_lead="source_doc / source_page:")

    # ---------- 6. INGESTION PIPELINE ----------
    doc.add_heading("6. Ingestion Pipeline", level=1)
    add_para(
        doc,
        "The pipeline is config-driven, idempotent, and resume-safe. New sources are "
        "added by editing config/sources.yaml — never by changing code. Each stage has "
        "an explicit data contract (see backend/app/ingestion/contracts.py)."
    )
    add_table_from_grid(doc, ["Stage", "Input", "Output", "Notes"], [
        ["1. Fetch",   "Source URL / manual drop", "FetchedItem (raw bytes + checksum)", "Polite, rate-limited HTTP or manual-drop"],
        ["2. Extract", "FetchedItem", "Raw text + table JSON",                            "pdfplumber for prose, camelot for rate tables"],
        ["3. Parse",   "Raw text",      "Ordered ParsedUnit stream",                      "Structure-aware: respects Section / sub-section / proviso / Explanation"],
        ["4. Chunk",   "ParsedUnit",    "Chunk records (parent + child)",                 "Adds breadcrumb prefix; no blind windowing"],
        ["5. Embed",   "Chunk.text",    "vector(1024)",                                   "bge-m3, EMBED_MAX_LENGTH cap for CPU efficiency"],
        ["6. Index",   "Chunk + embedding", "Postgres rows + HNSW + GIN",                 "Both dense (HNSW) and sparse (FTS) indexes"],
    ])
    add_para(
        doc,
        "Idempotency is guaranteed by the source checksum: re-running the pipeline on "
        "unchanged sources is a no-op. Celery beat re-runs the pipeline nightly to "
        "pick up new CBDT circulars; only new checksums are re-embedded."
    )

    # ---------- 7. RETRIEVAL & QUERY FLOW ----------
    doc.add_heading("7. Retrieval & Query Flow", level=1)
    add_para(
        doc,
        "When a user submits a question, the system runs ten stages in sequence. The "
        "key mental model: embeddings find which provision applies; the LLM only "
        "phrases the answer. If retrieval is poor, the LLM cannot recover."
    )
    flow_rows = [
        ["1",  "Embed query",            "bge-m3 produces a 1024-dim vector representing the question."],
        ["2",  "Dense vector search",    "pgvector HNSW returns ~20 nearest chunks by cosine distance."],
        ["3",  "Sparse FTS search",      "Postgres GIN tsvector returns ~20 keyword-matched chunks (in parallel with step 2)."],
        ["4",  "Merge (RRF)",            "Reciprocal Rank Fusion combines both candidate lists into ~30 unique hits."],
        ["5",  "Cross-encoder rerank",   "bge-reranker-v2-m3 scores each (query, chunk) pair and keeps top-K (typically 5)."],
        ["6",  "Parent expansion",       "Narrow child hits are swapped for their full parent Section for LLM context."],
        ["7",  "Grounding gate",         "If the top rerank score is below threshold, refuse and return — the LLM is not called."],
        ["8",  "Prompt assembly",        "System + statute context + chat history + user question, in cache-friendly order."],
        ["9",  "LLM generation",         "vLLM with batching; streaming response by default."],
        ["10", "Post-process & return",  "Parse citations from the answer, verify each cited section appears in retrieved chunks, log to audit trail."],
    ]
    add_table_from_grid(doc, ["#", "Stage", "What happens"], flow_rows)

    # ---------- 8. CHAT HISTORY ----------
    doc.add_heading("8. Chat History & Conversational Memory", level=1)
    add_para(
        doc,
        "For multi-turn conversations, the system maintains a hybrid memory model "
        "combining recency, semantic recall, and the fresh statute retrieval that "
        "happens on every turn."
    )
    doc.add_heading("8.1 The three tiers", level=2)
    add_bullet(doc, "Always included verbatim. Provides conversational coherence.", bold_lead="Last 5 messages:")
    add_bullet(doc, "Older user messages are embedded and stored per (user_id, chat_id). Once the chat exceeds a token threshold (not a fixed message count), the top 2–3 semantically relevant prior turns are retrieved and inserted as context.", bold_lead="Semantic recall over chat history:")
    add_bullet(doc, "Independent retrieval over the law corpus on every turn — never trusts the chat history to remember law correctly.", bold_lead="Fresh statute retrieval:")

    doc.add_heading("8.2 Critical design rules", level=2)
    add_bullet(doc, "Statute chunks and chat messages live in different tables. Different access control, different recency weighting, different retrieval logic.", bold_lead="Separate vector index from the statute corpus.")
    add_bullet(doc, "Embedding past assistant outputs perpetuates any previous hallucinations. Embed user questions only; store the assistant response as metadata for display.", bold_lead="Embed user questions, not assistant answers.")
    add_bullet(doc, "“5 messages” may be 200 tokens or 8000 tokens. Gate on total token count of conversation history.", bold_lead="Trigger by token count, not message count.")

    # ---------- 9. EVALUATION METHODOLOGY ----------
    doc.add_heading("9. Evaluation Methodology", level=1)
    add_para(
        doc,
        "Picking a chunking strategy or an LLM by intuition is the most common reason "
        "RAG products underperform. The honest method is an ablation study against a "
        "fixed evaluation set."
    )

    doc.add_heading("9.1 The eval set", level=2)
    add_para(
        doc,
        "A hand-curated set of approximately 30–50 questions, each paired with the "
        "ground-truth Section or Rule citation that should appear in a correct answer. "
        "Coverage should include: easy single-section lookups, multi-section reasoning, "
        "rate-table queries, recently-amended provisions, and edge cases that require "
        "refusal. The set is reused forever and acts as a regression test every time "
        "the embedding model, LLM, or chunking strategy changes."
    )

    doc.add_heading("9.2 The three metrics, measured independently", level=2)
    add_table_from_grid(doc, ["Metric", "What it measures", "Cost"], [
        ["Recall@k",            "Did the correct section appear in top-K retrieved chunks?", "Free — string match"],
        ["Citation correctness","Did the LLM cite the right section in its answer?",         "Free — regex"],
        ["Answer fidelity",     "Is every factual claim grounded in retrieved chunks?",      "Expensive — LLM-judge or human"],
    ])
    add_para(
        doc,
        "Always start with Recall@k. If retrieval misses the right section, nothing "
        "downstream matters — and Recall@k can be measured on hundreds of questions "
        "in minutes without invoking the LLM at all. Most teams burn money on answer-"
        "quality evaluation when the actual bottleneck is retrieval."
    )

    doc.add_heading("9.3 Ablation grid", level=2)
    add_para(
        doc,
        "Change one knob at a time, hold the rest fixed. The five high-leverage knobs are:"
    )
    add_bullet(doc, "structure-aware vs fixed-512 vs fixed-1024 tokens.",        bold_lead="Chunk granularity:")
    add_bullet(doc, "3 vs 5 vs 10.",                                              bold_lead="Top-K retrieved:")
    add_bullet(doc, "on / off.",                                                  bold_lead="Cross-encoder reranker:")
    add_bullet(doc, "on / off.",                                                  bold_lead="Parent expansion:")
    add_bullet(doc, "on / off.",                                                  bold_lead="Breadcrumb prefix:")
    add_para(doc, "Set LLM temperature to 0 during evaluation to remove sampling noise from configuration noise.")

    # ---------- 10. TOKEN EFFICIENCY ----------
    doc.add_heading("10. Token Efficiency at Inference", level=1)

    doc.add_heading("10.1 Where the tokens actually go per turn", level=2)
    add_code_block(doc, (
        "[system prompt + refusal rules]    ~  200–500    tok   (fixed, cacheable)\n"
        "[retrieved statute chunks]         ~ 1500–4000   tok   ← biggest controllable\n"
        "[chat history (last N + recall)]   ~  500–2000   tok   ← grows unbounded if ignored\n"
        "[user prompt]                      ~   50–200    tok\n"
        "[LLM output]                       ~  200–600    tok   ← you set the cap\n"
        "─────────────────────────────────────────────────\n"
        "                                    ~ 2500–7000+ tok per turn"
    ))

    doc.add_heading("10.2 High-leverage moves, in order", level=2)
    add_bullet(doc, "Place all stable content (system, persona, refusal rules) at the very start of the prompt. vLLM, Anthropic, and OpenAI all support prefix caching — repeat turns pay ~10% of the original input cost.", bold_lead="Prompt prefix caching:")
    add_bullet(doc, "“Always retrieve 10” is wasteful. Easy lookups need K=3; comparative questions need K=8. A cheap classifier or simple heuristic picks K per turn.", bold_lead="Dynamic top-K:")
    add_bullet(doc, "Expand to the full parent Section only when the child is short or the user query is broad. For narrow factual questions, the child alone suffices.", bold_lead="Conditional parent expansion:")
    add_bullet(doc, "Beyond message N, replace older turns with a regenerated 200-token rolling summary. Keeps chat history bounded.", bold_lead="Rolling-summary chat history:")
    add_bullet(doc, "Tax answers rarely need more than 400 output tokens. Combined with a strict format instruction (“≤3 sentences + citation”), this cuts both cost and decode latency.", bold_lead="max_tokens cap:")

    doc.add_heading("10.3 Medium-leverage moves", level=2)
    add_bullet(doc, "The verbose breadcrumb is great for embedding but heavy in the prompt. Compress “Income Tax Act 1961 > Chapter VIA > Section 80C > sub-section (2) > clause (a)” to “ITA 1961 s.80C(2)(a)” when feeding the LLM.", bold_lead="Compress breadcrumb at LLM-time:")
    add_bullet(doc, "If retrieval returns both a parent Section and its child sub-section, drop the redundant child to avoid paying twice for overlapping text.", bold_lead="Deduplicate retrieved chunks:")
    add_bullet(doc, "Some Sections are very long. If a retrieved chunk exceeds ~800 tokens, truncate around the matched span — recall has already happened.", bold_lead="Truncate long Sections at the LLM boundary:")

    doc.add_heading("10.4 The trap", level=2)
    add_para(
        doc,
        "Never measure tokens in isolation. A configuration that uses 30% fewer tokens "
        "but fails to retrieve the right Section 20% more often is a regression, not "
        "a win. Always tie token cost to Recall@k from the eval set."
    )

    # ---------- 11. HARDWARE REQUIREMENTS ----------
    doc.add_heading("11. Hardware Requirements", level=1)

    doc.add_heading("11.1 Where VRAM and RAM are consumed", level=2)
    add_code_block(doc, (
        "VRAM consumers:                                  System RAM consumers:\n"
        "─────────────────                                ──────────────────────\n"
        "bge-m3 embedding model       ~2.3 GB             Postgres + indexes (hot)  4–8 GB\n"
        "bge-reranker-v2-m3           ~2.3 GB             Redis cache               0.5–2 GB\n"
        "LLM weights                  varies              MinIO (raw files)         0.5 GB\n"
        "LLM KV cache (batch×ctx)     grows fast          API + worker + beat       2–3 GB\n"
        "CUDA / torch overhead        ~2 GB               OS + page cache buffer    4–8 GB"
    ))

    doc.add_heading("11.2 Four deployment tiers", level=2)
    add_table_from_grid(
        doc,
        ["Tier", "GPU", "LLM", "System RAM", "CPU", "Concurrent users", "Notes"],
        [
            ["Tier 0 — Dev",      "None (CPU)",                      "Qwen 2.5 7B Q4 (llama.cpp)",      "16 GB",     "8 cores",  "1, ~3–5 tok/s",  "Laptop development only"],
            ["Tier 1 — Small",    "1× 24 GB (L4 / RTX 4090 / A10G)", "Qwen 2.5 14B AWQ-Int4",           "32 GB",     "8 cores",  "20–50",          "Recommended starting point"],
            ["Tier 2 — Medium",   "1× A100 40 GB or 2× L4",          "Qwen 14B FP16 or Llama 70B Int4", "64 GB",     "16 cores", "100–200",        "Split DB onto its own node"],
            ["Tier 3 — Large",    "2–4× A100 80 GB or 1–2× H100",    "Llama 3.3 70B FP16",              "128–256 GB","32+ cores","500–1000+",      "Bottleneck shifts to orchestration"],
        ],
    )

    doc.add_heading("11.3 Choosing the tier honestly", level=2)
    add_bullet(doc, "50 officers logged in is not 50 concurrent — real concurrency is typically 10–20% of logged-in count.", bold_lead="Peak concurrent users (not registered users):")
    add_bullet(doc, "“First token in <1 s” forces vLLM + GPU + warm models. “Full answer in <5 s” is comfortable on Tier 1.", bold_lead="p95 latency target:")
    add_bullet(doc, "If answers must stay in India, self-host is mandatory (Tier 1 floor). If not, retrieval-only local + LLM-via-API often beats Tier 1 on quality.", bold_lead="Data-residency constraints:")

    # ---------- 12. DEPLOYMENT PATTERNS ----------
    doc.add_heading("12. Deployment Patterns", level=1)

    doc.add_heading("12.1 Pattern A — Single box (dev / very small prod)", level=2)
    add_para(
        doc,
        "All services in one Docker Compose stack on one VM. Simplest to operate. "
        "Limited by the single GPU's VRAM. Suitable for development and pilot "
        "deployments with under 20 concurrent users."
    )

    doc.add_heading("12.2 Pattern B — Two-VM split (recommended production)", level=2)
    add_para(
        doc,
        "The LLM lives on a dedicated GPU VM. The web application, retrieval pipeline, "
        "database, and embedding / reranker services live on a separate CPU VM. The two "
        "communicate over private VPC networking. This is the standard production "
        "topology because the LLM is the only workload that requires a GPU; everything "
        "else runs comfortably on cheaper CPU hardware."
    )
    add_code_block(doc, (
        "┌──────────────────────────────────────┐         ┌────────────────────────────┐\n"
        "│   VM 1 — App / Retrieval (CPU)       │  HTTPS  │   VM 2 — LLM (GPU)         │\n"
        "│  ──────────────────────────────────  │ ──────► │  ───────────────────────   │\n"
        "│  • Frontend (React/nginx)            │ private │  • vLLM serving Qwen/Llama │\n"
        "│  • API (FastAPI)                     │  VPC    │  • Loads model in VRAM     │\n"
        "│  • Worker + Beat (Celery)            │         │  • Exposes /v1/completions │\n"
        "│  • Postgres + pgvector               │ ◄────── │                            │\n"
        "│  • Redis                             │         │  GPU: 24 GB (L4 / 4090)    │\n"
        "│  • MinIO                             │         │  RAM: 16 GB                │\n"
        "│  • ml-server (embeddings + reranker) │         │  CPU: 4 cores              │\n"
        "│                                      │         │  Disk: 100 GB              │\n"
        "│  GPU: none  ·  RAM: 32 GB            │         │  No public IP              │\n"
        "│  CPU: 8 cores  ·  Disk: 200 GB NVMe  │         └────────────────────────────┘\n"
        "└──────────────────────────────────────┘"
    ))

    doc.add_heading("12.3 Configuration changes for the split", level=2)
    add_para(doc, "On VM 1 (the application box), update .env:")
    add_code_block(doc, (
        "LLM_BACKEND=vllm\n"
        "LLM_API_BASE=http://10.0.1.5:8000     # ← private IP of VM 2\n"
        "LLM_MODEL_NAME=Qwen/Qwen2.5-14B-Instruct\n"
        "LLM_API_KEY=<shared-secret>"
    ))
    add_para(doc, "On VM 2 (the GPU box), run only the vLLM container:")
    add_code_block(doc, (
        "docker run --gpus all -p 8000:8000 \\\n"
        "  -v ~/hfcache:/root/.cache/huggingface \\\n"
        "  vllm/vllm-openai:latest \\\n"
        "  --model Qwen/Qwen2.5-14B-Instruct \\\n"
        "  --quantization awq \\\n"
        "  --max-model-len 8192 \\\n"
        "  --api-key $LLM_API_KEY"
    ))
    add_para(
        doc,
        "No application code changes are needed because the LLM is already abstracted "
        "behind the LLMClient interface in backend/app/services/llm.py."
    )

    doc.add_heading("12.4 Network and security rules", level=2)
    add_bullet(doc, "Both VMs in the same VPC and same availability zone (cross-AZ adds ~1–2 ms; cross-region adds 50–200 ms).")
    add_bullet(doc, "VM 2 must have no public IP. Only VM 1 may reach VM 2 on port 8000.")
    add_bullet(doc, "Authentication between VMs: API key in header is sufficient for VPC-internal traffic. Add mTLS only if compliance demands it.")
    add_bullet(doc, "VM 2 ingress: allow tcp/8000 from VM 1 private IP only; allow tcp/22 from bastion only; deny everything else.")

    doc.add_heading("12.5 New failure modes to handle", level=2)
    add_table_from_grid(doc, ["Scenario", "Mitigation"], [
        ["VM 2 (LLM) down",        "Health-check + graceful refusal; retrieved chunks still surfaced for manual review"],
        ["Transient network blip", "Exponential backoff retries (3 tries: 100 / 300 / 900 ms)"],
        ["VM 2 restarted",         "Models reload (~30–60 s); add a warm-up cron after restart"],
        ["VM 2 overloaded",        "Client-side timeout at 10 s, return clear error to the user"],
    ])

    # ---------- 13. PERFORMANCE OPTIMIZATION ----------
    doc.add_heading("13. Performance Optimization", level=1)

    doc.add_heading("13.1 Realistic latency budget", level=2)
    add_table_from_grid(doc, ["Stage", "Realistic", "Floor"], [
        ["Query embedding (bge-m3, GPU)",       "10–20 ms",  "5 ms"],
        ["Vector search (HNSW, current scale)", "2–10 ms",   "<2 ms"],
        ["Sparse search (FTS)",                 "5–20 ms",   "3 ms"],
        ["Reranker (30 pairs, GPU)",            "30–80 ms",  "15 ms"],
        ["Parent expansion + DB fetch",         "5–15 ms",   "3 ms"],
        ["Total retrieval",                     "50–150 ms", "~30 ms"],
        ["LLM generation (unavoidable cost)",   "1000–5000 ms (streamed)", "depends on model"],
    ])
    add_para(
        doc,
        "Sub-100 ms is realistic for the retrieval portion. The LLM portion is "
        "irreducibly in the seconds — the right UX move is token-by-token streaming, "
        "so the user sees output appearing within ~300–800 ms of pressing Enter."
    )

    doc.add_heading("13.2 Highest-leverage optimisations", level=2)
    add_bullet(doc, "First /ask after restart currently takes ~30–60 s because models lazy-load. A warm-up call on container start fixes this. Free 1000× speedup on cold-start.", bold_lead="Add a warm-up hook:")
    add_bullet(doc, "bge-m3 on CPU is ~50–100 ms; on GPU it is ~5–15 ms. The embedding and reranker should be on the same GPU as the LLM (they fit comfortably in 5 GB alongside).", bold_lead="Move embedding and reranker to GPU:")
    add_bullet(doc, "Dense and sparse retrieval should run in parallel, not sequentially. Saves 10–30 ms per query.", bold_lead="Parallelise dense + sparse search:")
    add_bullet(doc, "If the top dense score is above 0.85, skip the reranker entirely. Easy queries no longer pay the 30–80 ms reranker cost.", bold_lead="Confidence-based reranker skip:")
    add_bullet(doc, "Cache hash(normalized_query) → embedding in Redis with a 1-hour TTL. Cache hit is 1 ms instead of 15 ms. Hit rate is often 30–50% for FAQ-style traffic.", bold_lead="Query-embedding cache:")

    doc.add_heading("13.3 Concurrency, not just single-user latency", level=2)
    add_bullet(doc, "The ml-server should batch incoming embedding and rerank requests onto the GPU (gather 16 calls, run one kernel). 5–10× throughput improvement at moderate load.", bold_lead="Batch ML-server requests:")
    add_bullet(doc, "ml-server, api, and worker are all stateless — replicate behind a load balancer. Postgres remains single-master with read replicas.", bold_lead="Horizontal replication of stateless services:")
    add_bullet(doc, "Watch queue depth and p95 latency, not CPU%. CPU lags reality on GPU-bound workloads.", bold_lead="Autoscaling signals:")

    # ---------- 14. SECURITY ----------
    doc.add_heading("14. Security, Compliance & Refusal Logic", level=1)

    doc.add_heading("14.1 The refusal guarantee", level=2)
    add_para(
        doc,
        "The grounding gate in backend/app/services/rag.py is the central trust "
        "mechanism. If the top reranker score falls below threshold, the LLM is never "
        "called — the system refuses with a clear message. This produces a measurable "
        "false-negative rate (some valid questions get refused), which is the right "
        "tradeoff: a refusal can be retried; a confidently wrong tax answer cannot be "
        "undone."
    )

    doc.add_heading("14.2 Citation verification", level=2)
    add_para(
        doc,
        "After the LLM generates an answer, citations are parsed and matched against "
        "the retrieved chunks. If the LLM cites a section that was not in its context, "
        "the citation is rejected and flagged in the audit log. This catches the "
        "residual hallucination cases that slip past the grounding gate."
    )

    doc.add_heading("14.3 Tenant isolation", level=2)
    add_bullet(doc, "Per-department audit log of every query, retrieval result, and answer.")
    add_bullet(doc, "Seat-based licensing prevents account sharing beyond licensed seats.")
    add_bullet(doc, "All chat history queries scoped to (user_id, chat_id); cross-tenant retrieval impossible by construction.")

    doc.add_heading("14.4 Data residency", level=2)
    add_para(
        doc,
        "For deployments that require Indian data residency, all VMs must be in an "
        "Indian region (AWS Mumbai ap-south-1, GCP Mumbai asia-south1, or on-premises). "
        "In such deployments, the LLM cannot be replaced with a foreign API (Anthropic, "
        "OpenAI), and self-hosting on Tier 1 hardware is the minimum viable option."
    )

    # ---------- 15. FUTURE ROADMAP ----------
    doc.add_heading("15. Future Roadmap", level=1)
    add_para(
        doc,
        "The current architecture supports several capability extensions without "
        "structural change. Each is independent and can be prioritised based on user "
        "demand."
    )
    add_table_from_grid(doc, ["Initiative", "Effort", "Value"], [
        ["Deterministic tax calculator (compute_tax tool)",          "Medium", "Removes a class of LLM arithmetic errors; very high"],
        ["Live web search integration (Tavily / Brave, whitelisted)", "Medium", "Coverage of post-cutoff circulars and news"],
        ["Case-law ingestion (ITAT / HC / SC)",                       "Large",  "Enables “has this been litigated?” queries"],
        ["Query router (statute / web / calc / clarify)",             "Small",  "Halves latency and cost on simple queries"],
        ["Multi-step agent loop (ReAct-style)",                       "Large",  "Quality lift on cross-section reasoning"],
        ["GST and Customs domains",                                   "Small",  "Already supported by the domain axis; YAML-only"],
        ["MS Word add-in for tax drafting",                           "Medium", "High value for practitioner workflows"],
    ])

    # ---------- 16. APPENDIX ----------
    doc.add_page_break()
    doc.add_heading("16. Appendix — Example Chunk Record (full)", level=1)
    add_para(
        doc,
        "Reproduced from §5.1 for reference. This is the canonical shape of one row "
        "in the chunks table after the full ingestion pipeline."
    )
    add_code_block(doc, EXAMPLE_CHUNK_JSON)

    add_para(doc, "")
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("— End of Document —")
    r.italic = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(out_path)


EXAMPLE_CHUNK_JSON = """{
  "chunk_id": "ita1961_s80E_ss1",

  "text": "Income Tax Act 1961 > Chapter VIA > Section 80E > sub-section (1)\\n\\nIn computing the total income of an assessee, being an individual, there shall be deducted, in accordance with and subject to the provisions of this section, any amount paid by him in the previous year, out of his income chargeable to tax, by way of interest on loan taken by him from any financial institution or any approved charitable institution for the purpose of pursuing his higher education or for the purpose of higher education of his relative.",

  "body": "In computing the total income of an assessee, being an individual, there shall be deducted ... interest on loan taken by him from any financial institution or any approved charitable institution for the purpose of pursuing his higher education or for the purpose of higher education of his relative.",

  "breadcrumb": "Income Tax Act 1961 > Chapter VIA > Section 80E > sub-section (1)",

  "level": "subsection",

  "act_name": "Income Tax Act 1961",
  "chapter": "VIA",
  "section_number": "80E",
  "subsection": "(1)",
  "clause": null,
  "proviso_no": null,
  "explanation_no": null,
  "rule_number": null,
  "subrule": null,

  "effective_date": "2006-04-01",
  "status": "in_force",
  "applicable_assessment_years": ["2006-07", "current"],

  "amended_by": [
    {"act": "Finance Act 2006", "year": 2006, "change": "inserted"},
    {"act": "Finance Act 2009", "year": 2009, "change": "scope expanded to include relative"}
  ],

  "cross_references": [
    {"section": "10(23C)",   "context": "referenced for 'approved charitable institution'"},
    {"section": "80G(2)(a)", "context": "alternative path for charitable institution definition"}
  ],

  "content_type": "prose",
  "domain": "income_tax",
  "jurisdiction": "IN",

  "parent_chunk_id": "ita1961_s80E",
  "parent_index": 0,

  "source_doc": "Income-tax-Act-1961_2026_2026-06-18_02-47-37_b9326c_en.pdf",
  "source_page": 643,
  "source_checksum": "b9326c...",
  "ingested_at": "2026-06-23T22:15:00Z",
  "schema_version": "1.0",

  "embedding": "[0.0214, -0.1387, 0.0892, ..., 0.0341]   # 1024-dim from bge-m3"
}"""


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parent.parent
    out_path = project_root / "BharathTax_Documentation.docx"
    build_doc(out_path)
    print(f"Wrote {out_path}  ({out_path.stat().st_size / 1024:.1f} KB)")
