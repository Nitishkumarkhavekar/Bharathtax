"""Generate a professional introduction letter for BharatTax as a .docx.

Plain Times New Roman throughout, no design flourishes. The BharatTax logo
sits at the top-left corner. All addresses / contact fields are left as
[placeholders] for the marketing team to fill in.

Run: python scripts/generate_intro_letter.py
Output: BharatTax-Introduction-Letter.docx (project root).
"""
from __future__ import annotations

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _force_font(style, name: str) -> None:
    """python-docx forgets to set the East-Asian / complex-script font when
    you only set style.font.name — meaning a Word install could substitute
    Calibri in some paragraphs. Set every rFonts attribute explicitly so
    the whole document renders in the requested face."""
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), name)


def build() -> Document:
    doc = Document()

    # ---- Base font: Times New Roman 11 pt ---------------------------------
    _force_font(doc.styles["Normal"], "Times New Roman")
    doc.styles["Normal"].font.size = Pt(11)

    # ---- Page margins (standard A4 letter) --------------------------------
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    # ---- Logo (top-left) --------------------------------------------------
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.add_run().add_picture("docs/Bharathtax-logo.png", width=Inches(1.8))

    # ---- Date -------------------------------------------------------------
    doc.add_paragraph()  # blank spacer
    doc.add_paragraph("[Date]")

    # ---- Salutation -------------------------------------------------------
    doc.add_paragraph("Respected Sir / Madam,")

    # ---- Subject line (bold) ---------------------------------------------
    subj_p = doc.add_paragraph()
    subj_r = subj_p.add_run(
        "Subject: Introduction to BharatTax — an AI-powered research, "
        "reading and drafting platform built for the Indian income-tax "
        "ecosystem."
    )
    subj_r.bold = True

    # ---- Body paragraphs --------------------------------------------------
    # Tuple entries render the first element bold (lead-in phrase) and the
    # second element in regular weight, on the same line — matches the
    # "First, Lexlegis research is grounded in legal sources." pattern.
    body: list = [
        (
            "I take the liberty of writing to introduce ",
            "BharatTax",
            ", a grounded and reliable artificial-intelligence platform "
            "built specifically for Indian income-tax practice. BharatTax "
            "combines a curated corpus of verified Indian tax sources "
            "— the Income-tax Act, 1961, the Income-tax Rules, 1962, "
            "the CBDT circulars and notifications, GST notifications, and "
            "Ministry of Corporate Affairs filings — with a live "
            "retrieval layer over judicial precedents from the Supreme "
            "Court, the High Courts, and the Income Tax Appellate "
            "Tribunal. Today, the platform serves as a trusted research, "
            "reading and drafting assistant to Chartered Accountants, "
            "Assessing Officers, Commissioner (Appeals) benches, corporate "
            "finance and secretarial teams, tax counsel, and taxpayer-side "
            "professionals across the country."
        ),
        (
            "Unlike general-purpose artificial-intelligence systems and "
            "GPTs, BharatTax operates upon architectural principles "
            "specifically designed for tax reliability and accuracy."
        ),
        (
            "First, BharatTax research is grounded in primary law. ",
            None,
            "Every response is anchored to verified Indian tax statutes, "
            "rules, circulars, and judicial precedents. Rather than relying "
            "on statistical language models alone, the platform retrieves "
            "and reasons over actual legal documents from its curated "
            "database. Every claim in the answer is footnoted to the exact "
            "section, sub-clause, rule, notification or reported judgment "
            "it stands on, and the source is one click away from the reader."
        ),
        (
            "Second, BharatTax reads the document, not the OCR. ",
            None,
            "Users may upload sale deeds (English and bilingual), "
            "assessment notices, show-cause letters, appellate orders, and "
            "any other tax-related document. The platform extracts facts "
            "verbatim, flags what is missing, and answers only from what "
            "is actually present in the file — never from what a "
            "general model would assume ought to be there. Where a "
            "critical figure is absent — for example, the Stamp Duty "
            "Value in a sale deed for a Section 50C determination — "
            "BharatTax refuses to reverse-engineer the value and instead "
            "requests the missing document, protecting the professional "
            "from a demonstrably incorrect conclusion."
        ),
        (
            "Third, BharatTax is right-shaped for who is asking. ",
            None,
            "The same underlying evidence produces different depth for "
            "different audiences: a Chartered Accountant receives a "
            "computation framework with conditional language and defence "
            "points; an Assessing Officer receives verification points and "
            "a suggested questionnaire; a Company Secretary or Chief "
            "Financial Officer receives an executive verdict with a risk "
            "rating and an actionable compliance checklist; and an "
            "individual taxpayer or founder receives a plain-English brief "
            "of what to do next and what documents to preserve. The "
            "reasoning discipline is identical; the format is tuned to the "
            "reader."
        ),
        (
            "Fourth, every answer is self-audited before it reaches the user. ",
            None,
            "A post-generation review pass detects reverse-engineered tax "
            "figures, initials-only or otherwise fabricated case citations, "
            "over-confident legal conclusions where a material fact is "
            "missing, and other patterns that would compromise professional "
            "reliance. When such a pattern is detected, a visible warning "
            "is prepended to the answer before the reader sees it. There "
            "are no silent failures."
        ),
        (
            "Fifth, BharatTax drafts appellate orders through a "
            "structured six-module pipeline. ",
            None,
            "For Commissioner (Appeals) benches and Assessing Officers "
            "preparing orders, the platform generates Facts, Deficiencies, "
            "Scope, Compliance, Findings, and the Order itself — each "
            "module individually cited, individually editable, and exported "
            "to a signable Microsoft Word document ready to be placed in "
            "front of a bench."
        ),
        (
            "The platform is deployed with tenant-scoped data isolation, "
            "audit-logged access, and role-based section allowlists — "
            "control lives with the firm, wing, or department that adopts "
            "it. No document uploaded by a user is used to train shared "
            "models."
        ),
        (
            "We would welcome the opportunity to demonstrate BharatTax to "
            "your team, to arrange a trial for your officers or partners, "
            "and to discuss how the platform can be tailored to your "
            "specific workflow. Please write to us at the address, e-mail "
            "or telephone number cited below to schedule a demonstration."
        ),
        (
            "Thank you for your time and consideration."
        ),
    ]

    for entry in body:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if isinstance(entry, tuple):
            for i, chunk in enumerate(entry):
                if chunk is None:
                    continue
                run = p.add_run(chunk)
                # First element of a bold-lead tuple gets bold; also the
                # 3-item pattern (lead, name, rest) bolds the middle name.
                if i == 0 and len(entry) >= 3:
                    run.bold = True
                elif len(entry) == 3 and i == 1:
                    run.bold = True
        else:
            p.add_run(entry)

    # ---- Sign-off ---------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("[Name]")
    doc.add_paragraph("[Designation]")
    doc.add_paragraph("For BharatTax")

    # ---- Footer block (plain, no design) ---------------------------------
    doc.add_paragraph()
    footer_lines = [
        "[Company Legal Name Private Limited]",
        "CIN: [Corporate Identification Number]",
        "Registered Office: [Full postal address, City — PIN, State, India]",
        "Website: [website]   |   E-mail: [contact e-mail]   |   Telephone: [contact number]",
    ]
    for line in footer_lines:
        fp = doc.add_paragraph(line)
        fp.paragraph_format.space_after = Pt(2)
        for r in fp.runs:
            r.font.size = Pt(10)

    return doc


if __name__ == "__main__":
    out_path = "BharatTax-Introduction-Letter.docx"
    build().save(out_path)
    print(f"Saved: {out_path}")
